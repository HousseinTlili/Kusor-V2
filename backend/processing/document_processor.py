# backend/processing/document_processor.py
"""
DocumentProcessor — handles the full lifecycle of turning uploaded files
(PDF, DOCX, TXT) into searchable chunks stored in PostgreSQL, ChromaDB,
and Neo4j, while triggering obligation extraction and temporal graph building.
"""

from __future__ import annotations

import hashlib
import io
import logging
import os
import re
import uuid
from datetime import datetime, date, timezone
from typing import List, Optional, Tuple, Dict, Any

import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract

from backend.extensions import db, get_chroma_collection, get_neo4j_manager
from backend.models.chunk import Chunk
from backend.models.document import Document
from backend.processing.obligation_extractor import ObligationExtractor

logger = logging.getLogger(__name__)

# Regex patterns for structural segmentation
_TITRE_RE = re.compile(
    r"^\s*(TITRE|Titre)\s+([IVXLCDM]+|PREMIER|DEUXIÈME|TROISIÈME|QUATRIÈME|CINQUIÈME)",
    re.MULTILINE,
)
_CHAPITRE_RE = re.compile(
    r"^\s*(CHAPITRE|Chapitre)\s+([IVXLCDM]+|\d+|PREMIER|DEUXIÈME|TROISIÈME|QUATRIÈME|CINQUIÈME)",
    re.MULTILINE,
)
_SECTION_RE = re.compile(
    r"^\s*(SECTION|Section)\s+(\d+|[IVXLCDM]+|PREMIÈRE|DEUXIÈME|TROISIÈME)",
    re.MULTILINE,
)
_ARTICLE_RE = re.compile(
    r"^\s*(Article|ARTICLE)\s+(\d+|premier|Premier|PREMIER)",
    re.MULTILINE,
)
_CIRCULAR_REF_RE = re.compile(
    r"[Cc]irculaire\s+[Nn]°?\s*(\d{4}-\d{1,2})", re.IGNORECASE
)
_DATE_RE = re.compile(
    r"(\d{1,2}(?:er)?)\s+(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(\d{4})",
    re.IGNORECASE,
)

_MONTH_MAP = {
    "janvier": 1, "février": 2, "mars": 3, "avril": 4,
    "mai": 5, "juin": 6, "juillet": 7, "août": 8,
    "septembre": 9, "octobre": 10, "novembre": 11, "décembre": 12,
}


class DocumentProcessor:
    """
    Orchestrates complete document ingestion:
    1. Text extraction (PyMuPDF with Tesseract OCR fallback for PDF, python-docx for DOCX, UTF-8 for TXT)
    2. Structural segmentation (Titre / Chapitre / Section / Article)
    3. Overlapping text chunking
    4. Database persistence (PostgreSQL + ChromaDB with v3 metadata)
    5. Temporal Knowledge Graph + Obligation Graph construction
    """

    def __init__(
        self,
        chroma_collection: Any = None,
        neo4j_manager: Any = None,
        ollama_base_url: str = "http://localhost:11434",
        llm_model: str = "qwen2.5:7b",
        chunk_size: int = 800,
        chunk_overlap: int = 100,
    ):
        self._chroma = chroma_collection or get_chroma_collection()
        self._neo4j = neo4j_manager or get_neo4j_manager()
        self._ollama_base_url = ollama_base_url
        self._llm_model = llm_model
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._obligation_extractor = ObligationExtractor(
            ollama_base_url=ollama_base_url,
            llm_model=llm_model,
        )

    def process_document(
        self,
        filepath_or_stream: Any,
        doc_id: Optional[str] = None,
        circular_ref: Optional[str] = None,
        title: Optional[str] = None,
        doc_type: str = "circular",
        filename: Optional[str] = None,
    ) -> Document:
        """Process file path or stream into text and index chunks/graph."""
        actual_filename = filename
        if isinstance(filepath_or_stream, str) and os.path.exists(filepath_or_stream):
            actual_filename = os.path.basename(filepath_or_stream)
            if actual_filename.endswith(".pdf"):
                raw_text = self._extract_pdf(filepath_or_stream)
            elif actual_filename.endswith(".docx"):
                raw_text = self._extract_docx(filepath_or_stream)
            else:
                raw_text = self._extract_txt(filepath_or_stream)
        else:
            raw_text = str(filepath_or_stream)

        # Detect circular reference from filename if like 2018-12.pdf
        effective_ref = circular_ref
        if not effective_ref and actual_filename:
            ref_match = re.match(r"^(\d{4}-\d{1,2})", actual_filename)
            if ref_match:
                effective_ref = ref_match.group(1)

        doc_title = title
        if not doc_title:
            if effective_ref:
                doc_title = f"Circulaire BCT N° {effective_ref}"
            elif actual_filename:
                doc_title = actual_filename
            else:
                doc_title = "Document"

        return self.process_text_content(
            raw_text,
            title=doc_title,
            doc_type=doc_type,
            circular_ref=effective_ref,
            filename=actual_filename,
            doc_id=doc_id,
        )

    def process_text_content(
        self,
        raw_text: str,
        title: str,
        doc_type: str = "circular",
        existing_doc: Optional[Document] = None,
        circular_ref: Optional[str] = None,
        filename: Optional[str] = None,
        doc_id: Optional[str] = None,
    ) -> Document:
        """Process raw text directly without file upload."""
        if not raw_text.strip():
            raise ValueError("No text provided")

        if not circular_ref:
            circular_ref = self._extract_circular_reference(raw_text)
        date_issued = self._extract_date(raw_text)
        content_hash = hashlib.sha256(raw_text.encode("utf-8")).hexdigest()

        if not existing_doc:
            if doc_id:
                existing_doc = Document.query.get(doc_id)
            if not existing_doc and circular_ref:
                existing_doc = Document.query.filter(
                    (Document.number == circular_ref) | (Document.circular_reference == circular_ref)
                ).first()
            if not existing_doc and filename:
                existing_doc = Document.query.filter(Document.filename == filename).first()

        if existing_doc:
            doc = existing_doc
            doc.title = title
            doc.filename = filename or doc.filename
            doc.doc_type = doc_type
            doc.number = circular_ref or doc.number
            doc.circular_reference = circular_ref or doc.circular_reference
            doc.date_issued = date_issued or doc.date_issued
            doc.content_hash = content_hash
            doc.raw_text = raw_text
            doc.indexation_state = "PROCESSING"
            doc.updated_at = datetime.now(timezone.utc)
            # Delete old chunks
            Chunk.query.filter_by(document_id=doc.id).delete()
        else:
            doc = Document(
                id=doc_id or str(uuid.uuid4()),
                title=title,
                filename=filename,
                doc_type=doc_type,
                number=circular_ref,
                circular_reference=circular_ref,
                date_issued=date_issued,
                content_hash=content_hash,
                raw_text=raw_text,
                indexation_state="PROCESSING",
                source="BCT Portal",
            )
            db.session.add(doc)

        db.session.commit()

        sections = self._segment_text(raw_text)
        chunks = self._create_chunks(sections, doc.id)

        for chunk in chunks:
            db.session.add(chunk)
        db.session.commit()

        if self._chroma:
            self._store_chromadb(chunks, doc)

        self._build_graph_and_obligations(doc, raw_text, sections)

        doc.indexation_state = "INDEXED"
        db.session.commit()

        try:
            if self._neo4j and (doc.circular_reference or doc.number):
                from backend.agent.propagation_agent import ChangePropagationAgent
                prop_agent = ChangePropagationAgent(self._neo4j)
                prop_agent.analyze_impact(doc.circular_reference or doc.number, document_id=doc.id)
        except Exception as e:
            logger.warning("Auto change propagation trigger failed: %s", e)

        logger.info("Successfully processed document %s (%d chunks)", doc.id, len(chunks))
        return doc


    # ── Text Extraction Methods ─────────────────────────────────

    @staticmethod
    def _extract_pdf(filepath: str) -> str:
        """Extract text from PDF using PyMuPDF; fallback to Tesseract OCR if text length < 50 chars."""
        doc = fitz.open(filepath)
        pages_text = []
        for page in doc:
            text = page.get_text()
            if len(text.strip()) < 50:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(img, lang="fra")
            pages_text.append(text)
        doc.close()
        return "\n".join(pages_text)

    @staticmethod
    def _extract_docx(filepath: str) -> str:
        """Extract text from DOCX document."""
        document = docx.Document(filepath)
        return "\n".join([p.text for p in document.paragraphs if p.text.strip()])

    @staticmethod
    def _extract_txt(filepath: str) -> str:
        """Extract text from UTF-8 encoded text file."""
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _extract_circular_reference(self, text: str) -> Optional[str]:
        match = _CIRCULAR_REF_RE.search(text)
        return match.group(1) if match else None

    def _extract_date(self, text: str) -> Optional[date]:
        match = _DATE_RE.search(text)
        if not match:
            return None
        day = int(match.group(1).replace("er", ""))
        month = _MONTH_MAP.get(match.group(2).lower())
        year = int(match.group(3))
        if month:
            try:
                return date(year, month, day)
            except ValueError:
                return None
        return None

    def _segment_text(self, text: str) -> List[Tuple[str, str]]:
        markers = []
        for pattern, prefix in [
            (_TITRE_RE, "Titre"),
            (_CHAPITRE_RE, "Chapitre"),
            (_SECTION_RE, "Section"),
            (_ARTICLE_RE, "Article"),
        ]:
            for m in pattern.finditer(text):
                markers.append((m.start(), f"{prefix} {m.group(2)}", m.end()))

        if not markers:
            return [("Document complet", text)]

        markers.sort(key=lambda x: x[0])
        sections = []
        for i, (start, title, content_start) in enumerate(markers):
            end = markers[i + 1][0] if i + 1 < len(markers) else len(text)
            content = text[content_start:end].strip()
            if content:
                sections.append((title, content))

        if markers[0][0] > 0:
            preamble = text[: markers[0][0]].strip()
            if preamble:
                sections.insert(0, ("Préambule", preamble))

        return sections

    def _create_chunks(
        self, sections: List[Tuple[str, str]], doc_id: str
    ) -> List[Chunk]:
        chunks = []
        idx = 0
        for section_title, section_content in sections:
            words = section_content.split()
            start = 0
            while start < len(words):
                end = min(start + self._chunk_size, len(words))
                chunk_text = " ".join(words[start:end])

                c = Chunk(
                    id=f"{doc_id}_{idx}",
                    document_id=doc_id,
                    chunk_index=idx,
                    content=chunk_text,
                    token_count=len(words[start:end]),
                    section_title=section_title,
                )
                chunks.append(c)
                idx += 1
                start += self._chunk_size - self._chunk_overlap
                if start >= len(words):
                    break
        return chunks

    def _store_chromadb(self, chunks: List[Chunk], doc: Document) -> None:
        if not self._chroma:
            return

        ids = [c.id for c in chunks]
        documents = [c.content for c in chunks]
        metadatas = [
            {
                "document_id": doc.id,
                "title": doc.title or "",
                "circular_reference": doc.circular_reference or "",
                "doc_type": doc.doc_type,
                "section_title": c.section_title or "",
                "chunk_index": c.chunk_index,
            }
            for c in chunks
        ]
        try:
            self._chroma.add(ids=ids, documents=documents, metadatas=metadatas)
        except Exception as e:
            logger.error("Failed to add vectors to ChromaDB: %s", e)

    def _build_graph_and_obligations(
        self, doc: Document, text: str, sections: List[Tuple[str, str]]
    ) -> None:
        if not self._neo4j:
            return

        title = doc.title or f"Circulaire BCT {doc.circular_reference or doc.id}"
        c_query = """
        MERGE (c:Circular {id: $id})
        SET c.title = $title,
            c.number = $ref,
            c.date_issued = $date,
            c.status = 'ACTIVE'
        """
        self._neo4j.run_query(
            c_query,
            {
                "id": doc.id,
                "title": title,
                "ref": doc.circular_reference,
                "date": doc.date_issued.isoformat() if doc.date_issued else None,
            },
        )

        obligations = self._obligation_extractor.extract_obligations(doc, sections, use_llm=False)
        for ob in obligations:
            ob_query = """
            MATCH (c:Circular {id: $doc_id})
            MERGE (o:Obligation {id: $ob_id})
            SET o.text = $text,
                o.obligation_type = $type
            MERGE (c)-[r:INTRODUCES]->(o)
            """
            self._neo4j.run_query(
                ob_query,
                {
                    "doc_id": doc.id,
                    "ob_id": ob.id,
                    "text": ob.text,
                    "type": ob.obligation_type,
                },
            )
